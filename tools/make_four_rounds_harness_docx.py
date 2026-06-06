from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "harness_four_rounds_assets"
OUT = DOCS / "耿同学agent_四轮harness流程图.docx"

BLUE = "#14528F"
BLUE2 = "#2F7FC8"
FILL = "#EEF5FC"
WARN = "#FFF5E6"
DARK = "#1F2933"
GRAY = "#5B6775"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        try:
            return ImageFont.truetype(item, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=8)
    return box[2] - box[0], box[3] - box[1]


def box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    subtitle: str = "",
    fill: str = FILL,
    outline: str = BLUE2,
) -> None:
    draw.rounded_rectangle(xy, radius=22, fill=fill, outline=outline, width=4)
    x1, y1, x2, y2 = xy
    title_font = font(30, True)
    sub_font = font(22)
    if subtitle:
        content = title + "\n" + subtitle
        w, h = text_size(draw, content, title_font)
        draw.multiline_text(
            ((x1 + x2) / 2, y1 + 28),
            title,
            fill=BLUE,
            font=title_font,
            anchor="ma",
            align="center",
            spacing=10,
        )
        draw.multiline_text(
            ((x1 + x2) / 2, y1 + 78),
            subtitle,
            fill=DARK,
            font=sub_font,
            anchor="ma",
            align="center",
            spacing=8,
        )
    else:
        w, h = text_size(draw, title, title_font)
        draw.multiline_text(
            ((x1 + x2) / 2, (y1 + y2) / 2 - h / 2),
            title,
            fill=BLUE,
            font=title_font,
            anchor="ma",
            align="center",
            spacing=8,
        )


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str = "") -> None:
    draw.line((start, end), fill=BLUE, width=5)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 11), (ex - 18 * direction, ey + 11)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 11, ey - 18 * direction), (ex + 11, ey - 18 * direction)]
    draw.polygon(points, fill=BLUE)
    if label:
        mx, my = (sx + ex) // 2, (sy + ey) // 2
        draw.rounded_rectangle((mx - 82, my - 23, mx + 82, my + 23), radius=12, fill="white", outline="#C8D6E5")
        draw.text((mx, my - 13), label, font=font(19), fill=GRAY, anchor="ma")


def header(draw: ImageDraw.ImageDraw, title: str, subtitle: str) -> None:
    draw.text((70, 45), title, font=font(42, True), fill=BLUE)
    draw.text((70, 100), subtitle, font=font(24), fill=GRAY)
    draw.line((70, 138, 1530, 138), fill="#D8E4F0", width=3)


def footer(draw: ImageDraw.ImageDraw, note: str) -> None:
    draw.rounded_rectangle((70, 835, 1530, 895), radius=18, fill="#F8FAFC", outline="#D7E0EA", width=2)
    draw.text((95, 853), note, font=font(22), fill=GRAY)


def new_canvas() -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1600, 930), "white")
    return img, ImageDraw.Draw(img)


def round1(path: Path) -> Path:
    img, d = new_canvas()
    header(d, "第一轮：从论文里抽工程事实", "目标：把 PDF 变成可以审查的 facts JSON")
    box(d, (80, 250, 310, 380), "论文 PDF", "公式 / 图表 / 实验段")
    box(d, (420, 250, 650, 380), "文本切块", "paper_chunks")
    box(d, (760, 250, 1010, 380), "LLM 读取", "只抽工程事实")
    box(d, (1120, 250, 1390, 380), "本地校验", "schema + chunk 回指")
    box(d, (575, 560, 875, 700), "facts JSON", "通过后进入第二轮")
    box(d, (1060, 560, 1360, 700), "本地 fallback", "多次失败时兜底", WARN, "#E6A23C")
    arrow(d, (310, 315), (420, 315))
    arrow(d, (650, 315), (760, 315))
    arrow(d, (1010, 315), (1120, 315))
    arrow(d, (1255, 380), (760, 560), "通过")
    arrow(d, (1255, 380), (1210, 560), "失败过多")
    footer(d, "这一轮不写代码，只回答：论文到底声称用了哪些模型、参数、指标和图表。")
    img.save(path)
    return path


def round2(path: Path) -> Path:
    img, d = new_canvas()
    header(d, "第二轮：把事实变成复现任务", "目标：明确要复现哪张图、哪条曲线、哪些输出列")
    box(d, (85, 235, 345, 370), "facts JSON", "第一轮产物")
    box(d, (85, 475, 345, 610), "论文 chunks", "补充上下文")
    box(d, (500, 350, 780, 490), "LLM 规划任务", "实验 / 指标 / baseline")
    box(d, (930, 235, 1220, 370), "tasks JSON", "任务清单")
    box(d, (930, 475, 1220, 610), "index JSON", "图表地图")
    box(d, (1310, 350, 1535, 490), "本地校验", "引用 / 口径")
    arrow(d, (345, 300), (500, 405))
    arrow(d, (345, 540), (500, 435))
    arrow(d, (780, 405), (930, 300))
    arrow(d, (780, 435), (930, 540))
    arrow(d, (1220, 300), (1310, 405))
    arrow(d, (1220, 540), (1310, 435))
    footer(d, "这一轮回答：要跑什么实验，以及本地输出怎样才算和论文结果能对比。")
    img.save(path)
    return path


def round3(path: Path) -> Path:
    img, d = new_canvas()
    header(d, "第三轮：生成并运行复现项目", "目标：LLM 写代码，但代码必须经过本地护栏")
    box(d, (70, 215, 285, 345), "任务输入", "facts + tasks")
    box(d, (410, 215, 670, 345), "分文件生成", "每文件 <= 100 行")
    box(d, (800, 215, 1035, 345), "代码审查", "路径 / import / 依赖")
    box(d, (1165, 215, 1435, 345), "落地项目", "repro_project")
    box(d, (410, 505, 670, 635), "smoke 运行", "先小规模试跑")
    box(d, (800, 505, 1035, 635), "full config", "完整配置运行")
    box(d, (1165, 505, 1435, 635), "outputs", "CSV / PNG / summary")
    box(d, (615, 690, 895, 810), "失败修复", "LLM repair / OpenHands")
    arrow(d, (285, 280), (410, 280))
    arrow(d, (670, 280), (800, 280))
    arrow(d, (1035, 280), (1165, 280))
    arrow(d, (1300, 345), (540, 505))
    arrow(d, (670, 570), (800, 570))
    arrow(d, (1035, 570), (1165, 570))
    arrow(d, (920, 635), (755, 690), "失败")
    arrow(d, (755, 690), (540, 635), "重跑")
    footer(d, "这一轮不是相信代码能跑，而是用本地规则、依赖检查、smoke 和 full 运行连续筛。")
    img.save(path)
    return path


def round4(path: Path) -> Path:
    img, d = new_canvas()
    header(d, "第四轮：审查复现结果并生成报告", "目标：拿本地结果回头对比论文结论")
    box(d, (80, 205, 330, 335), "本地 outputs", "数据 / 图像 / summary")
    box(d, (80, 455, 330, 585), "论文证据", "相关页面图 + chunks")
    box(d, (495, 330, 775, 470), "多模态 LLM", "看图 + 看表 + 看上下文")
    box(d, (925, 205, 1210, 335), "result_review", "逐实验可信度")
    box(d, (925, 455, 1210, 585), "risk_report", "多维风险汇总")
    box(d, (1325, 330, 1530, 470), "Word 报告", "review.docx")
    arrow(d, (330, 270), (495, 385))
    arrow(d, (330, 520), (495, 415))
    arrow(d, (775, 385), (925, 270))
    arrow(d, (775, 415), (925, 520))
    arrow(d, (1210, 270), (1325, 385))
    arrow(d, (1210, 520), (1325, 415))
    footer(d, "这一轮只给出复现可信度、差异原因和人工复核建议，不直接判定论文造假。")
    img.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor.from_string("14528F")


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("耿同学agent harness 四轮流程图")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("14528F")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("一页一轮：从论文输入，到代码运行，再到结果审查")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("5B6775")


def add_round(doc: Document, title: str, image_path: Path, bullets: list[str], page_break: bool = True) -> None:
    if page_break:
        doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(7.25))
    table = doc.add_table(rows=len(bullets), cols=1)
    table.autofit = False
    for row, bullet in zip(table.rows, bullets):
        cell = row.cells[0]
        set_cell_shading(cell, "F8FAFC")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(bullet)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("1F2933")


def build_doc() -> Path:
    DOCS.mkdir(exist_ok=True)
    ASSETS.mkdir(exist_ok=True)
    paths = [
        round1(ASSETS / "round_01_facts.png"),
        round2(ASSETS / "round_02_tasks.png"),
        round3(ASSETS / "round_03_codegen_runtime.png"),
        round4(ASSETS / "round_04_result_review.png"),
    ]
    doc = Document()
    set_doc_style(doc)
    add_title(doc)
    doc.add_paragraph(
        "这份文档只保留最少文字，把耿同学agent的 harness 路径拆成四张图。"
        "新手可以按 1 -> 2 -> 3 -> 4 的顺序理解整个系统。"
    )
    add_round(
        doc,
        "第一轮：论文 -> 工程事实",
        paths[0],
        [
            "输入是 PDF；输出是 engineering_facts.json。",
            "重点不是复现，而是先把论文里的模型、参数、指标、图表抽清楚。",
            "如果 LLM 多次给不出合格 JSON，会启用本地 fallback，保证流程能继续但风险升高。",
        ],
    )
    add_round(
        doc,
        "第二轮：工程事实 -> 复现任务",
        paths[1],
        [
            "输入是 facts JSON 和论文上下文；输出是 repro_tasks.json 与 experiment_index.json。",
            "这一轮决定到底复现哪张图、哪条曲线、哪个指标。",
            "本地校验会检查任务是否引用真实事实、指标口径是否足够明确。",
        ],
    )
    add_round(
        doc,
        "第三轮：复现任务 -> 本地 outputs",
        paths[2],
        [
            "LLM 可以自由写代码，但每个文件最多 100 行，并且要通过依赖、import、安全和路径审查。",
            "运行顺序改成先 smoke，再跑 full config；最终产物必须是 CSV、PNG 和 summary JSON。",
            "失败时进入修复链：LLM repair、OpenHands 候选副本、本地模板 fallback。",
        ],
    )
    add_round(
        doc,
        "第四轮：本地 outputs -> 审查报告",
        paths[3],
        [
            "把本地数据、图像、summary、论文页面图和任务元数据一起交给多模态 LLM。",
            "输出 result_review 与 risk_report，评价每个实验的复现可信度和差异原因。",
            "最后生成 review.docx / result_review.docx，但结论只表达风险，不直接判定论文造假。",
        ],
    )
    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
