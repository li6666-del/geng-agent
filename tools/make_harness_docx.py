from __future__ import annotations

import os
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "harness_architecture_assets"
DOCX_PATH = OUT_DIR / "耿同学agent_harness架构说明.docx"

FONT = "Microsoft YaHei"
ACCENT = "2E74B5"
DARK = "1F4D78"
MUTED = "666666"
FILL = "E8EEF5"
LIGHT = "F6F8FB"


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)
    diagrams = build_diagrams()

    doc = Document()
    setup_document(doc)
    add_title(doc)
    add_quick_map(doc)
    add_section(doc, "1. 这套 harness 到底是什么")
    add_para(
        doc,
        "这里的 harness 可以理解成“把论文复现审查这件事串起来的工作台”。"
        "它不是单个 LLM，也不是单个脚本，而是一套负责调度、校验、运行、返修和出报告的流水线。"
        "LLM 负责理解论文和生成候选内容，本地代码负责把关、执行、重试、修复和兜底。"
    )
    add_callout(
        doc,
        "一句话版本",
        "耿同学agent = PDF 论文输入 + LLM 结构化理解 + 本地安全运行 + 结果审查 + Word 报告输出。",
    )

    add_section(doc, "2. 总流程图")
    add_image(doc, diagrams["flow"], "图 1  从 PDF 到 Word 审查报告的主链路")
    add_para(
        doc,
        "从新手角度看，只需要记住三个阶段：先把论文变成结构化事实，再把事实变成复现实验和代码，"
        "最后本地执行并把结果反向送回 LLM 做审查。中间每一步都会留下 JSON、日志或报告，便于追踪问题。"
    )

    add_section(doc, "3. harness 分层架构")
    add_image(doc, diagrams["layers"], "图 2  耿同学agent 的分层结构")
    add_layer_table(doc)

    add_section(doc, "4. 四个核心 JSON：系统的骨架")
    add_para(
        doc,
        "新手最容易迷糊的是：为什么这么多 JSON？可以把它们看作流水线之间的交接单。"
        "LLM 可以读懂 JSON，本地程序也可以严格校验 JSON，所以 JSON 是人、模型、程序三方都能对齐的中间语言。"
    )
    add_json_table(doc)

    add_section(doc, "5. 第三轮代码生成：保留自由发挥，但有护栏")
    add_para(
        doc,
        "当前第三轮不是完全模板化，也不是完全放飞。LLM 仍然可以写 run_experiment.py 和 src/ 下的实现逻辑，"
        "但必须在固定项目结构、固定输出产物、安全规则、依赖规则、200 行文件上限内完成。"
    )
    add_image(doc, diagrams["codegen"], "图 3  第三轮代码生成和本地审查的关系")
    add_codegen_table(doc)

    add_section(doc, "6. 本地执行与失败恢复闭环")
    add_image(doc, diagrams["recovery"], "图 4  smoke/full 运行、LLM repair 和 fallback")
    add_para(
        doc,
        "smoke 是小规模试跑，full 是完整 config 运行。现在 smoke 通过不会直接进入最终审查，"
        "必须继续跑完整 config。若运行失败，先尝试 LLM repair；"
        "仍不可靠时才使用本地确定性模板 fallback。"
    )

    add_section(doc, "7. 自动化闭环：为什么叫 agent")
    add_agent_loop_table(doc)
    add_para(
        doc,
        "当前系统没有单独的监督命令；agent 性主要体现在主流水线内部的闭环。"
        "它会读取 schema 校验、audit、runtime_result、repair_logs 和 result_review，再决定重试、补抽、修复、兜底或停止。"
    )

    add_section(doc, "8. 输出目录怎么读")
    add_output_tree(doc)
    add_para(
        doc,
        "如果你只想看结论，先打开 review.docx 和 result_review.docx。"
        "如果你想定位问题，按 audit -> runtime_result.json -> repair_logs -> risk_report.json 的顺序看。"
    )

    add_section(doc, "9. 新手读项目的推荐路线")
    add_numbered_steps(
        doc,
        [
            "先看 README.md，理解命令入口和输出目录。",
            "再看 case_xxx/paper_chunks.json，确认论文被拆成了哪些文本块。",
            "看 engineering_facts.json 和 repro_tasks.json，确认 LLM 抽了什么、打算复现什么。",
            "看 experiment_index.json，把任务对应回论文图表和页码。",
            "看 repro_project/，理解 LLM 生成的可执行项目。",
            "看 runtime_result.json，判断 smoke/full 是否真正跑过。",
            "最后看 review.docx、result_review.docx 和 risk_report.json。",
        ],
    )

    add_section(doc, "10. 这套架构的边界")
    add_callout(
        doc,
        "重要声明",
        "系统输出的是复现风险、差异分析和人工复核建议，不直接判定论文造假。"
        "如果论文没有公开数据或关键参数，系统会诚实记录假设和局限，而不是假装完全复现。",
    )

    add_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    set_east_asia(normal, FONT)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        set_east_asia(style, FONT)


def set_east_asia(style_or_run, font_name: str) -> None:
    element = style_or_run._element
    if hasattr(element, "get_or_add_rPr"):
        rpr = element.get_or_add_rPr()
    else:
        rpr = element.rPr
    rfonts = rpr.get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(key), font_name)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("耿同学agent Harness 架构说明")
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(DARK)
    set_east_asia(run, FONT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("给新手看的论文复现审查流水线地图")
    run.font.name = FONT
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia(run, FONT)


def add_quick_map(doc: Document) -> None:
    rows = [
        ("输入", "论文 PDF/TXT/Markdown"),
        ("核心过程", "论文解析 -> 事实抽取 -> 复现任务 -> 代码生成 -> 本地执行 -> 结果审查"),
        ("关键护栏", "JSON schema、路径校验、安全扫描、依赖审查、smoke/full 双阶段运行"),
        ("输出", "review.docx、result_review.docx、risk_report.json、repro_project/outputs"),
    ]
    add_kv_table(doc, rows)


def add_section(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)


def add_para(doc: Document, text: str) -> None:
    for chunk in textwrap.wrap(text, width=78):
        p = doc.add_paragraph(chunk)
        p.paragraph_format.space_after = Pt(6)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    set_east_asia(r, FONT)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)


def add_image(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.runs[0]
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1800, 7560])
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "说明"
    for cell in hdr:
        shade_cell(cell, FILL)
        bold_cell(cell)
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    polish_table(table)


def add_layer_table(doc: Document) -> None:
    rows = [
        ("用户入口层", "CLI: review/status/doctor；Web: /api/runs", "接收用户命令或上传 PDF，决定输出到哪个 case。"),
        ("Pipeline 层", "ReviewPipeline", "串起每一阶段：解析、LLM 调用、项目生成、运行、报告。"),
        ("LLM 层", "MiniMax/OpenAI-compatible API", "做论文理解、复现任务、代码生成、结果审查。"),
        ("本地校验层", "schema/security/outputs", "不信任 LLM，负责结构审查、安全扫描、依赖校验、产物校验。"),
        ("执行与修复层", "runner/LLM repair/fallback", "在隔离思路下运行代码，失败后尝试修复或兜底。"),
        ("报告层", "review/result_review/docx", "把过程证据整理成 Markdown、JSON 和 Word。"),
    ]
    add_matrix_table(doc, ["层", "代表模块", "小白解释"], rows, [1500, 2600, 5260])


def add_json_table(doc: Document) -> None:
    rows = [
        ("engineering_facts.json", "论文里的工程事实", "模型/信道/调制/指标/图表/baseline 等事实，必须回指 paper_chunks。"),
        ("repro_tasks.json", "要复现的实验任务", "每个任务说明目标图表、指标公式、输出列、预期趋势和依赖事实。"),
        ("experiment_index.json", "实验地图", "把复现任务对应回论文页码、图表和 chunk，方便审查定位。"),
        ("risk_report.json", "风险总账", "汇总信息完整性、实现忠实度、运行可靠性、结果贴合度和最终 verdict。"),
    ]
    add_matrix_table(doc, ["文件", "它是什么", "为什么重要"], rows, [2200, 2300, 4860])


def add_codegen_table(doc: Document) -> None:
    rows = [
        ("固定结构", "LLM 只能生成 README、requirements、config、run_experiment 和 src 下几个文件。"),
        ("200 行上限", "每个生成文件最多 200 行，避免一口气写出难修的大工程。"),
        ("依赖规则", "代码 import 的第三方库必须写进 requirements，且只能用当前环境允许的库。"),
        ("运行产物", "必须生成 CSV、PNG、summary JSON，不能只打印文字。"),
        ("本地审查", "路径、安全、依赖、语法、产物都会被本地代码硬检查。"),
    ]
    add_matrix_table(doc, ["护栏", "含义"], rows, [2200, 7160])


def add_agent_loop_table(doc: Document) -> None:
    rows = [
        ("JSON 重试", "LLM 输出解析或 schema 校验失败时，带错误信息重新请求。"),
        ("事实补抽", "图表或 figure claim 覆盖不足时，定向追加 facts gap round。"),
        ("任务补齐", "可复现实验没有任务覆盖时，定向追加 tasks gap round。"),
        ("代码修复", "受限运行失败时，触发 LLM repair。"),
        ("模板兜底", "生成项目无法通过本地校验或运行验收时，写入确定性 fallback。"),
        ("结果复审", "有真实输出时，把本地 PNG/CSV/论文页交给多模态模型逐实验审查。"),
    ]
    add_matrix_table(doc, ["闭环动作", "什么时候用"], rows, [2200, 7160])


def add_output_tree(doc: Document) -> None:
    lines = [
        "case_xxx/",
        "  paper_chunks.json",
        "  engineering_facts.json",
        "  repro_tasks.json",
        "  experiment_index.json",
        "  repro_project/",
        "    config.json / config_smoke.json / run_experiment.py / src/",
        "    outputs/results.csv / *.png / summary.json",
        "    repair_logs/",
        "  runtime_result.json",
        "  result_review.json / result_review.md / result_review.docx",
        "  risk_report.json",
        "  review.md / review.docx",
        "  audit/",
        "  reflections/",
    ]
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "FBFCFE")
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    p = cell.paragraphs[0]
    for line in lines:
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string("222222")
    p.paragraph_format.space_after = Pt(0)


def add_numbered_steps(doc: Document, steps: list[str]) -> None:
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(f"{i}. ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(DARK)
        p.add_run(step)


def add_matrix_table(doc: Document, headers: list[str], rows: list[tuple], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade_cell(cell, FILL)
        bold_cell(cell)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    polish_table(table)


def set_table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(sum(widths)))
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcW = cell._tc.get_or_add_tcPr().tcW
            tcW.type = "dxa"
            tcW.w = widths[idx]


def polish_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(10)
                    set_east_asia(r, FONT)


def bold_cell(cell) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(DARK)


def shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("耿同学agent harness 架构说明 - 复现风险评估，不直接判定论文造假")
    r.font.name = FONT
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia(r, FONT)


def build_diagrams() -> dict[str, Path]:
    return {
        "flow": draw_flow_diagram(ASSET_DIR / "01_flow.png"),
        "layers": draw_layers_diagram(ASSET_DIR / "02_layers.png"),
        "codegen": draw_codegen_diagram(ASSET_DIR / "03_codegen.png"),
        "recovery": draw_recovery_diagram(ASSET_DIR / "04_recovery.png"),
    }


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        if path and os.path.exists(path):
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def new_canvas(w=1600, h=900):
    return Image.new("RGB", (w, h), "white"), ImageDraw.Draw(Image.new("RGB", (1, 1)))


def draw_box(draw, xy, text, fill="F6F8FB", outline="2E74B5", text_color="12324A", radius=18, size=28):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill="#" + fill, outline="#" + outline, width=3)
    lines = wrap_text(draw, text, font(size, True), x2 - x1 - 32)
    total_h = len(lines) * (size + 8)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font(size, True))
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), line, fill="#" + text_color, font=font(size, True))
        y += size + 8


def arrow(draw, start, end, color="1F4D78"):
    draw.line([start, end], fill="#" + color, width=5)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    draw.polygon(pts, fill="#" + color)


def wrap_text(draw, text, fnt, max_w):
    out = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            if draw.textlength(line + ch, font=fnt) <= max_w:
                line += ch
            else:
                out.append(line)
                line = ch
        if line:
            out.append(line)
    return out


def draw_flow_diagram(path: Path) -> Path:
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "从论文到报告：主 harness 流程", fill="#" + DARK, font=font(40, True))
    boxes = [
        ((70, 130, 330, 250), "论文 PDF\nchunks JSON"),
        ((430, 130, 690, 250), "工程事实\nfacts JSON"),
        ((790, 130, 1050, 250), "复现任务\ntasks JSON"),
        ((1150, 130, 1480, 250), "实验地图\nindex JSON"),
        ((70, 390, 330, 510), "生成代码\n项目文件"),
        ((430, 390, 690, 510), "本地审查\n结构/安全"),
        ((790, 390, 1050, 510), "执行运行\nsmoke + full"),
        ((1150, 390, 1480, 510), "输出产物\nCSV/PNG/JSON"),
        ((430, 660, 690, 780), "结果审查\n二次审查"),
        ((790, 660, 1050, 780), "最终分级\nverdict"),
        ((1150, 660, 1480, 780), "Word 报告\nDOCX"),
    ]
    for xy, text in boxes:
        draw_box(d, xy, text)
    for s, e in [((330, 190), (430, 190)), ((690, 190), (790, 190)), ((1050, 190), (1150, 190)),
                 ((1315, 250), (1315, 390)), ((1150, 450), (1050, 450)), ((790, 450), (690, 450)),
                 ((430, 450), (330, 450)), ((200, 510), (560, 660)), ((690, 720), (790, 720)),
                 ((1050, 720), (1150, 720)), ((1315, 510), (560, 660))]:
        arrow(d, s, e)
    img.save(path)
    return path


def draw_layers_diagram(path: Path) -> Path:
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "分层架构：谁负责什么", fill="#" + DARK, font=font(40, True))
    layers = [
        ("用户入口层", "CLI: review / status / doctor；Web API", "E8EEF5"),
        ("配置与客户端层", "GENG_LLM_* / GENG_LLM2_* / GENG_GEN_* / OpenAI-compatible client", "F4F6F9"),
        ("主流水线层", "ReviewPipeline: facts -> tasks -> project -> runtime -> reports", "E8EEF5"),
        ("LLM 能力层", "MiniMax/OpenAI-compatible: 事实、任务、代码、结果审查", "F4F6F9"),
        ("本地护栏层", "Pydantic schema / security scan / requirements / output validation", "E8EEF5"),
        ("执行修复层", "runner / LLM repair / template fallback", "F4F6F9"),
        ("证据报告层", "audit / runtime_result / risk_report / review.docx", "E8EEF5"),
    ]
    y = 120
    for title, desc, fill in layers:
        d.rounded_rectangle((100, y, 1500, y + 82), radius=18, fill="#" + fill, outline="#" + ACCENT, width=3)
        d.text((140, y + 18), title, fill="#" + DARK, font=font(28, True))
        d.text((430, y + 20), desc, fill="#222222", font=font(24))
        y += 98
    img.save(path)
    return path


def draw_codegen_diagram(path: Path) -> Path:
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "第三轮：LLM 写代码，但在护栏里写", fill="#" + DARK, font=font(40, True))
    draw_box(d, (80, 160, 430, 310), "LLM 自由发挥\n写每个文件", "F6F8FB")
    draw_box(d, (620, 120, 980, 230), "固定项目结构\n9 个核心文件", "E8EEF5")
    draw_box(d, (620, 270, 980, 380), "每文件最多 200 行\n避免大工程失控", "E8EEF5")
    draw_box(d, (620, 420, 980, 530), "依赖必须写 requirements\n禁止乱 import", "E8EEF5")
    draw_box(d, (620, 570, 980, 680), "必须输出 CSV/PNG/summary\n可被本地审查", "E8EEF5")
    draw_box(d, (1160, 310, 1500, 500), "本地校验\n不通过就返修/兜底", "F4F6F9")
    for y in [175, 325, 475, 625]:
        arrow(d, (430, 235), (620, y))
    for y in [175, 325, 475, 625]:
        arrow(d, (980, y), (1160, 405))
    img.save(path)
    return path


def draw_recovery_diagram(path: Path) -> Path:
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "失败恢复：不中断，但会记录风险", fill="#" + DARK, font=font(40, True))
    boxes = [
        ((90, 150, 390, 270), "运行 smoke"),
        ((520, 150, 820, 270), "运行 full config"),
        ((950, 150, 1300, 270), "产物校验\nCSV/PNG/summary"),
        ((520, 410, 820, 530), "LLM repair\n返回修复 manifest"),
        ((520, 670, 820, 790), "template fallback\n本地确定性兜底"),
        ((950, 670, 1300, 790), "risk_report\n记录降级原因"),
    ]
    for xy, text in boxes:
        draw_box(d, xy, text)
    arrow(d, (390, 210), (520, 210))
    arrow(d, (820, 210), (950, 210))
    arrow(d, (670, 270), (670, 410))
    arrow(d, (670, 530), (670, 670))
    arrow(d, (820, 730), (950, 730))
    d.text((105, 315), "失败时进入修复链路；成功时继续下一步", fill="#555555", font=font(24))
    img.save(path)
    return path


if __name__ == "__main__":
    main()
