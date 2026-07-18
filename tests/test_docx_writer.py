from __future__ import annotations

import base64
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from geng_agent.docx_writer import (
    write_markdown_report_docx,
    write_result_review_docx,
    write_result_review_markdown_docx,
    write_review_docx,
)


TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


class DocxWriterTests(unittest.TestCase):
    def test_write_review_docx_creates_openable_report(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            path = root / "review.docx"

            write_review_docx(
                path,
                paper={"source_path": "paper.pdf", "chunk_count": 3},
                facts={
                    "paper_repro_type": "signal_chain",
                    "engineering_facts": [
                        {
                            "type": "channel_model",
                            "name": "AWGN",
                            "source": {"chunk_id": "text_c1"},
                            "confidence": "high",
                            "used_for_reproduction": True,
                        }
                    ],
                    "missing_information": [],
                },
                tasks={
                    "repro_tasks": [
                        {
                            "task_id": "reproduce_fig_1",
                            "target": "BER vs SNR",
                            "metric": "bit_error_rate",
                            "figure_or_claim": "Fig. 1",
                            "expected_trend": {"direction": "decreasing"},
                            "comparison": {"baselines": ["AWGN reference"]},
                        }
                    ]
                },
                risk_report={
                    "risk_level": "medium",
                    "missing_information_count": 0,
                    "assumptions_count": 1,
                    "risk_dimensions": {
                        "runtime_reliability": {"level": "medium", "evidence": ["runtime_enabled=False"]}
                    },
                    "findings": [],
                },
                validation={"required_files_present": True, "python_compiles": True},
                runtime_result={
                    "enabled": True,
                    "passed": True,
                    "attempts": [],
                    "requirements_warnings": [
                        {
                            "file": "tasks/demo.py",
                            "line": "1",
                            "message": "third-party import is not declared in requirements.txt: scipy.linalg (expected package scipy)",
                        }
                    ],
                },
                result_review_result={"enabled": False, "passed": None, "reason": "not run"},
                repro_project_dir=root / "repro_project",
            )

            document = Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_text = "\n".join(
                paragraph.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
            self.assertIn("耿同学agent 论文工程复现审查报告", text)
            self.assertIn("人工复核建议", text)
            self.assertIn("依赖告警", text)
            self.assertIn("scipy.linalg", table_text)
            self.assertGreaterEqual(len(document.tables), 5)

    def test_write_result_review_docx_creates_experiment_sections(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            path = root / "result_review.docx"

            write_result_review_docx(
                path,
                result_review={
                    "overall_result_credibility": "medium",
                    "overall_alignment": "partial_match",
                    "experiment_reviews": [
                            {
                                "task_id": "reproduce_fig_1",
                                "local_result_credibility": "medium",
                                "paper_alignment": "partial_match",
                                "scientific_verdict": "partially_supports_paper_claim",
                                "dimension_reviews": [
                                    {
                                        "dimension": dimension,
                                        "rating": "acceptable",
                                        "finding": f"{dimension} 维度有基础证据。",
                                        "evidence": ["outputs/results.csv"],
                                    }
                                    for dimension in [
                                        "artifact_coverage",
                                        "reproduction_logic",
                                        "trend_shape",
                                        "metric_axis_scale",
                                        "baseline_comparison",
                                        "statistical_reliability",
                                        "conclusion_support",
                                    ]
                                ],
                                "paper_result_summary": "Paper reports BER decreasing with SNR.",
                                "local_result_summary": "Local CSV shows BER decreasing.",
                                "differences": ["Only smoke-scale data were generated."],
                            "possible_causes": ["Fewer samples than the paper."],
                            "evidence": ["outputs/results.csv"],
                            "limitations": ["No precise figure digitization."],
                            "confidence": "medium",
                        }
                    ],
                    "cross_experiment_findings": ["Qualitative trend is present."],
                    "recommended_human_checks": ["Run full config.json."],
                    "note": "Risk review only.",
                },
                status={
                    "passed": True,
                    "result_review_path": str(root / "result_review.json"),
                    "result_review_markdown_path": str(root / "result_review.md"),
                },
            )

            document = Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("复现结果二次审查报告", text)
            self.assertIn("reproduce_fig_1", text)
            self.assertIn("多维审查", text)
            self.assertGreaterEqual(len(document.tables), 3)

    def test_write_result_review_markdown_docx_creates_openable_report(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            path = root / "result_review.docx"
            image_path = root / "local.png"
            paper_image_path = root / "paper.png"
            image_path.write_bytes(TINY_PNG)
            paper_image_path.write_bytes(TINY_PNG)

            write_result_review_markdown_docx(
                path,
                markdown_text=(
                    "## 1. reproduce_fig_1\n\n"
                    "### 图像对比\n\n"
                    "| 本地复现图 | 论文原图 |\n"
                    "|---|---|\n"
                    f"| ![本地复现图]({image_path}) | ![论文原图：Fig. 1]({paper_image_path}) |\n\n"
                    "### 简短审查结论\n\n"
                    "- 本地曲线趋势一致。\n"
                ),
                status={
                    "passed": True,
                    "mode": "codex_markdown_by_experiment",
                    "result_review_markdown_path": str(root / "result_review.md"),
                },
            )

            document = Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_text = "\n".join(
                paragraph.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
            self.assertIn("复现结果二次审查报告", text)
            self.assertIn("本地曲线趋势一致", text)
            self.assertIn("本地复现图", table_text)
            self.assertIn("论文原图：Fig. 1", table_text)
            self.assertNotIn("附录", text)
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(len(document.inline_shapes), 2)

    def test_generic_report_docx_resolves_relative_report_assets(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            asset = root / "report_assets" / "task_1" / "paper_target.png"
            asset.parent.mkdir(parents=True)
            asset.write_bytes(TINY_PNG)
            path = root / "reproduction_report.docx"

            write_markdown_report_docx(
                path,
                markdown_text="## task_1\n\n![论文原图](report_assets/task_1/paper_target.png)\n",
                title="本地复现报告",
                subtitle="参数与假设",
                base_dir=root,
            )

            document = Document(path)
            self.assertEqual(len(document.inline_shapes), 1)
            self.assertIn("本地复现报告", "\n".join(paragraph.text for paragraph in document.paragraphs))

    def test_write_result_review_markdown_docx_records_missing_images(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            path = root / "result_review.docx"
            missing = root / "missing.png"

            write_result_review_markdown_docx(
                path,
                markdown_text=f"## 1. reproduce_fig_1\n\n![missing figure]({missing})\n\nReviewer body.\n",
                status={"passed": True},
            )

            document = Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("图片缺失", text)
            self.assertIn(str(missing), text)
            self.assertEqual(len(document.inline_shapes), 0)

    def test_write_result_review_markdown_docx_records_missing_image_in_comparison_table(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            path = root / "result_review.docx"
            local_image = root / "local.png"
            missing = root / "missing.png"
            local_image.write_bytes(TINY_PNG)

            write_result_review_markdown_docx(
                path,
                markdown_text=(
                    "## 1. reproduce_fig_1\n\n"
                    "| 本地复现图 | 论文原图 |\n"
                    "|---|---|\n"
                    f"| ![本地复现图]({local_image}) | ![论文原图]({missing}) |\n"
                ),
                status={"passed": True},
            )

            document = Document(path)
            text = "\n".join(
                paragraph.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
            self.assertIn("图片缺失", text)
            self.assertIn(str(missing), text)
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(len(document.inline_shapes), 1)

    def test_write_result_review_markdown_docx_renders_multiple_images_in_one_table_cell(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            path = root / "result_review.docx"
            local_a = root / "local_a.png"
            local_b = root / "local_b.png"
            paper = root / "paper.png"
            for image_path in (local_a, local_b, paper):
                image_path.write_bytes(TINY_PNG)

            write_result_review_markdown_docx(
                path,
                markdown_text=(
                    "## 1. reproduce_fig_1\n\n"
                    "| 本地复现图 | 论文原图 |\n"
                    "|---|---|\n"
                    f"| ![本地复现图 A]({local_a}) <br /> ![本地复现图 B]({local_b}) "
                    f"| ![论文原图]({paper}) |\n"
                ),
                status={"passed": True},
            )

            document = Document(path)
            table_text = "\n".join(
                paragraph.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(len(document.inline_shapes), 3)
            self.assertIn("本地复现图 A", table_text)
            self.assertIn("本地复现图 B", table_text)
            self.assertNotIn("![", table_text)


if __name__ == "__main__":
    unittest.main()
