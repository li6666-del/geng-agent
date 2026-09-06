from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentObject

from .docx_styles import (
    _add_bullets,
    _add_heading,
    _add_kv_table,
    _add_labelled_bullets,
    _add_note,
    _add_table,
    _items_or_default,
    _safe_text,
)


RISK_LABELS = {
    "low": "低",
    "medium": "中",
    "high": "高",
    "unknown": "未知",
}
ALIGNMENT_LABELS = {
    "match": "匹配",
    "partial_match": "部分匹配",
    "mismatch": "不匹配",
    "inconclusive": "无法判断",
}
SCIENTIFIC_VERDICT_LABELS = {
    "supports_paper_claim": "支持论文主张",
    "partially_supports_paper_claim": "部分支持论文主张",
    "does_not_support_paper_claim": "不支持论文主张",
    "cannot_assess": "无法评估",
}
DIMENSION_RATING_LABELS = {
    "strong": "强",
    "acceptable": "可接受",
    "weak": "弱",
    "missing": "缺失",
    "unknown": "未知",
}
RESULT_REVIEW_DIMENSION_LABELS = {
    "artifact_coverage": "产物覆盖",
    "reproduction_logic": "复现逻辑",
    "trend_shape": "趋势走向",
    "metric_axis_scale": "指标与坐标轴",
    "baseline_comparison": "baseline 对比",
    "statistical_reliability": "统计可靠性",
    "conclusion_support": "结论支持度",
}


def populate_review_document(
    document: DocumentObject,
    *,
    paper: dict[str, Any],
    facts: dict[str, Any],
    tasks: dict[str, Any],
    risk_report: dict[str, Any],
    validation: dict[str, Any],
    runtime_result: dict[str, Any],
    result_review_result: dict[str, Any],
    repro_project_dir: Path,
) -> None:
    engineering_facts = facts.get("engineering_facts", [])
    missing = facts.get("missing_information", [])
    repro_tasks = tasks.get("repro_tasks", [])

    _add_heading(document, "一、基本信息", 1)
    _add_kv_table(
        document,
        [
            ("论文文件", paper.get("source_path")),
            ("文本块数量", paper.get("chunk_count")),
            ("论文复现类型", facts.get("paper_repro_type", "unknown")),
            ("复现项目目录", repro_project_dir),
            (
                "主风险等级",
                _label_with_raw(risk_report.get("risk_level"), RISK_LABELS),
            ),
        ],
    )

    _add_heading(document, "二、工程事实与缺失信息", 1)
    document.add_paragraph(
        f"本轮共抽取 {len(engineering_facts) if isinstance(engineering_facts, list) else 0} 条工程事实，"
        f"记录 {len(missing) if isinstance(missing, list) else 0} 条缺失信息。"
    )
    if isinstance(engineering_facts, list) and engineering_facts:
        rows = []
        for fact in engineering_facts[:12]:
            if not isinstance(fact, dict):
                continue
            source = (
                fact.get("source")
                if isinstance(fact.get("source"), dict)
                else {}
            )
            rows.append(
                [
                    fact.get("type", "unknown"),
                    fact.get("name", "unnamed"),
                    _label_with_raw(fact.get("confidence"), RISK_LABELS),
                    source.get("chunk_id", "unknown"),
                    "是" if fact.get("used_for_reproduction") else "否",
                ]
            )
        _add_table(
            document,
            ["类型", "名称", "置信度", "来源 chunk", "用于复现"],
            rows,
        )
    else:
        _add_note(document, "未抽取到可列出的工程事实。")

    if isinstance(missing, list) and missing:
        _add_heading(document, "缺失信息", 2)
        rows = []
        for item in missing:
            if isinstance(item, dict):
                rows.append(
                    [
                        item.get("name"),
                        item.get("why_needed"),
                        _label_with_raw(item.get("impact"), RISK_LABELS),
                    ]
                )
        _add_table(document, ["缺失项", "为什么需要", "影响"], rows)

    _add_heading(document, "三、复现任务", 1)
    if isinstance(repro_tasks, list) and repro_tasks:
        rows = []
        for task in repro_tasks:
            if not isinstance(task, dict):
                continue
            trend = (
                task.get("expected_trend")
                if isinstance(task.get("expected_trend"), dict)
                else {}
            )
            comparison = (
                task.get("comparison")
                if isinstance(task.get("comparison"), dict)
                else {}
            )
            rows.append(
                [
                    task.get("task_id"),
                    task.get("target"),
                    task.get("metric"),
                    task.get("figure_or_claim"),
                    trend.get("direction", "unknown"),
                    _join_items(comparison.get("baselines")),
                ]
            )
        _add_table(
            document,
            ["任务 ID", "目标", "指标", "图表/结论", "预期趋势", "基线"],
            rows,
        )
    else:
        _add_note(document, "未生成复现任务。")

    _add_heading(document, "四、风险摘要", 1)
    _add_kv_table(
        document,
        [
            (
                "总风险等级",
                _label_with_raw(risk_report.get("risk_level"), RISK_LABELS),
            ),
            ("缺失信息数量", risk_report.get("missing_information_count")),
            ("假设数量", risk_report.get("assumptions_count")),
            ("未解决任务证据缺口", risk_report.get("task_evidence_gap_count", 0)),
            ("判断边界", "复现风险评估，不等同于论文真伪判定"),
        ],
    )
    experiment_index = risk_report.get("experiment_index", {})
    experiments = (
        experiment_index.get("experiments")
        if isinstance(experiment_index, dict)
        else []
    )
    _add_heading(document, "Experiment Index", 1)
    if isinstance(experiments, list) and experiments:
        rows = []
        for experiment in experiments:
            if isinstance(experiment, dict):
                rows.append(
                    [
                        experiment.get("experiment_id"),
                        experiment.get("task_id"),
                        experiment.get("metric"),
                        experiment.get("figure_or_table"),
                        _join_items(experiment.get("limitations")),
                        _join_items(experiment.get("source_pages")),
                    ]
                )
        _add_table(
            document,
            [
                "Experiment",
                "Task",
                "Metric",
                "Figure/Table",
                "Evidence gaps",
                "Pages",
            ],
            rows,
        )
    else:
        _add_note(document, "experiment_index.json was not generated.")

    verdict = risk_report.get("reproducibility_verdict", {})
    _add_heading(document, "Final Reproducibility Verdict", 1)
    if isinstance(verdict, dict) and verdict:
        _add_kv_table(
            document,
            [
                ("verdict", verdict.get("verdict")),
                ("confidence", verdict.get("confidence")),
                ("recommended_action", verdict.get("recommended_action")),
            ],
        )
        _add_labelled_bullets(document, "Reasons", verdict.get("reasons"))
    else:
        _add_note(
            document,
            "Final reproducibility verdict was not generated.",
        )

    dimensions = risk_report.get("risk_dimensions", {})
    if isinstance(dimensions, dict) and dimensions:
        rows = []
        for name, dimension in dimensions.items():
            if isinstance(dimension, dict):
                rows.append(
                    [
                        name,
                        _label_with_raw(dimension.get("level"), RISK_LABELS),
                        _join_items(dimension.get("evidence")),
                    ]
                )
        _add_table(document, ["维度", "风险", "证据"], rows)

    findings = risk_report.get("findings", [])
    if isinstance(findings, list) and findings:
        _add_heading(document, "主要发现", 2)
        _add_bullets(
            document,
            [
                _safe_text(item.get("message", item))
                if isinstance(item, dict)
                else _safe_text(item)
                for item in findings
            ],
        )

    _add_heading(document, "五、本地运行与产物校验", 1)
    _add_kv_table(
        document,
        [
            ("必要文件齐全", validation.get("required_files_present")),
            ("Python 语法可编译", validation.get("python_compiles")),
            ("自动运行已启用", runtime_result.get("enabled")),
            ("自动运行是否通过", runtime_result.get("passed")),
            ("任务 full 运行覆盖", runtime_result.get("coverage", "未记录")),
            (
                "任务交付覆盖",
                runtime_result.get("delivery_coverage", "未记录"),
            ),
            (
                "主持人是否重复 full",
                runtime_result.get("host_repeated_full", False),
            ),
            (
                "依赖告警数",
                _runtime_requirement_warning_count(runtime_result),
            ),
        ],
    )
    requirement_warnings = _runtime_requirement_warnings(runtime_result)
    if requirement_warnings:
        _add_heading(document, "依赖告警", 2)
        rows = [
            [
                (
                    warning.get("file", "unknown")
                    if isinstance(warning, dict)
                    else "unknown"
                ),
                (
                    warning.get("line", "")
                    if isinstance(warning, dict)
                    else ""
                ),
                (
                    warning.get("message", warning)
                    if isinstance(warning, dict)
                    else warning
                ),
            ]
            for warning in requirement_warnings[:5]
        ]
        _add_table(document, ["文件", "行", "说明"], rows)
    artifacts = runtime_result.get("artifacts")
    if isinstance(artifacts, dict) and artifacts:
        rows = [[key, _safe_text(value)] for key, value in artifacts.items()]
        _add_table(document, ["产物检查项", "结果"], rows)

    _add_heading(document, "六、结果级二次审查状态", 1)
    _add_kv_table(
        document,
        [
            ("是否启用", result_review_result.get("enabled")),
            ("是否通过", result_review_result.get("passed")),
            (
                "结构化报告",
                result_review_result.get("result_review_path", "未生成"),
            ),
            (
                "可读报告",
                result_review_result.get(
                    "result_review_markdown_path",
                    "未生成",
                ),
            ),
            (
                "说明",
                result_review_result.get("reason")
                or result_review_result.get("error")
                or "无",
            ),
        ],
    )

    _add_heading(document, "七、人工复核建议", 1)
    _add_bullets(
        document,
        [
            "优先检查缺失信息和假设参数是否会改变核心结论。",
            "复现代码即使通过 smoke，也需要再运行完整 config.json 后比较论文图表。",
            "对 BER、吞吐量、准确率等指标，人工复核公式、统计口径、随机种子和样本规模。",
            "若结果级二次审查失败，需要确认多模态模型是否支持 image_url 输入。",
        ],
    )


def populate_result_review_document(
    document: DocumentObject,
    *,
    result_review: dict[str, Any],
    status: dict[str, Any] | None = None,
) -> None:
    status = status or {}
    _add_heading(document, "一、总体结论", 1)
    _add_kv_table(
        document,
        [
            (
                "本地结果总体可信度",
                _label_with_raw(
                    result_review.get("overall_result_credibility"),
                    RISK_LABELS,
                ),
            ),
            (
                "与原论文总体贴合度",
                _label_with_raw(
                    result_review.get("overall_alignment"),
                    ALIGNMENT_LABELS,
                ),
            ),
            ("审查状态", "通过" if status.get("passed") else "未记录"),
            ("结构化报告", status.get("result_review_path", "未记录")),
            (
                "Markdown 报告",
                status.get("result_review_markdown_path", "未记录"),
            ),
        ],
    )

    evidence = (
        status.get("evidence")
        if isinstance(status.get("evidence"), dict)
        else {}
    )
    if evidence:
        _add_heading(document, "输入证据摘要", 2)
        _add_kv_table(
            document,
            [
                ("CSV 摘要", _join_named_files(evidence.get("csv_summaries"))),
                ("summary JSON", _join_named_files(evidence.get("summary_json"))),
                ("本地图像", _join_named_files(evidence.get("output_images"))),
                ("论文页面图", _join_named_files(evidence.get("paper_images"))),
            ],
        )

    _add_heading(document, "二、逐实验点评", 1)
    experiment_reviews = result_review.get("experiment_reviews", [])
    if isinstance(experiment_reviews, list) and experiment_reviews:
        for index, review in enumerate(experiment_reviews, start=1):
            if not isinstance(review, dict):
                continue
            _add_heading(
                document,
                f"{index}. {review.get('task_id', 'unknown_task')}",
                2,
            )
            _add_table(
                document,
                ["项目", "内容"],
                [
                    [
                        "本地复现可信度",
                        _label_with_raw(
                            review.get("local_result_credibility"),
                            RISK_LABELS,
                        ),
                    ],
                    [
                        "论文贴合度",
                        _label_with_raw(
                            review.get("paper_alignment"),
                            ALIGNMENT_LABELS,
                        ),
                    ],
                    [
                        "科学结论支持",
                        _label_with_raw(
                            review.get("scientific_verdict"),
                            SCIENTIFIC_VERDICT_LABELS,
                        ),
                    ],
                    [
                        "判断置信度",
                        _label_with_raw(review.get("confidence"), RISK_LABELS),
                    ],
                    ["原论文结果摘要", review.get("paper_result_summary")],
                    ["本地结果摘要", review.get("local_result_summary")],
                ],
            )
            dimension_rows = []
            for dimension_review in _ordered_dimension_reviews(review):
                dimension = str(dimension_review.get("dimension", ""))
                dimension_rows.append(
                    [
                        _label_with_raw(
                            dimension,
                            RESULT_REVIEW_DIMENSION_LABELS,
                        ),
                        _label_with_raw(
                            dimension_review.get("rating"),
                            DIMENSION_RATING_LABELS,
                        ),
                        dimension_review.get("finding"),
                        _join_items(dimension_review.get("evidence")),
                    ]
                )
            if dimension_rows:
                _add_heading(document, "多维审查", 3)
                _add_table(
                    document,
                    ["维度", "评级", "判断", "证据"],
                    dimension_rows,
                )
            _add_labelled_bullets(
                document,
                "主要差异",
                review.get("differences"),
            )
            _add_labelled_bullets(
                document,
                "可能原因",
                review.get("possible_causes"),
            )
            _add_labelled_bullets(document, "证据", review.get("evidence"))
            _add_labelled_bullets(
                document,
                "局限性",
                review.get("limitations"),
            )
    else:
        _add_note(document, "未提供逐实验审查记录。")

    _add_heading(document, "三、跨实验观察", 1)
    _add_bullets(
        document,
        _items_or_default(
            result_review.get("cross_experiment_findings"),
            "未列出跨实验观察。",
        ),
    )

    _add_heading(document, "四、建议人工复核", 1)
    _add_bullets(
        document,
        _items_or_default(
            result_review.get("recommended_human_checks"),
            "未列出人工复核建议。",
        ),
    )

    _add_heading(document, "五、说明", 1)
    document.add_paragraph(_safe_text(result_review.get("note", "")))


def _runtime_requirement_warnings(runtime_result: dict[str, Any]) -> list[Any]:
    warnings = runtime_result.get("requirements_warnings")
    return warnings if isinstance(warnings, list) else []


def _runtime_requirement_warning_count(runtime_result: dict[str, Any]) -> int:
    return len(_runtime_requirement_warnings(runtime_result))


def _join_named_files(value: Any) -> str:
    if not isinstance(value, list):
        return _safe_text(value or "未记录")
    names = []
    for item in value:
        if isinstance(item, dict):
            names.append(
                _safe_text(
                    item.get("file")
                    or item.get("label")
                    or item.get("page")
                    or item
                )
            )
        else:
            names.append(_safe_text(item))
    return _join_items(names)


def _ordered_dimension_reviews(
    review: dict[str, Any],
) -> list[dict[str, Any]]:
    dimension_reviews = [
        item
        for item in review.get("dimension_reviews", [])
        if isinstance(item, dict)
    ]
    order = {
        dimension: index
        for index, dimension in enumerate(RESULT_REVIEW_DIMENSION_LABELS)
    }
    return sorted(
        dimension_reviews,
        key=lambda item: order.get(
            str(item.get("dimension", "")),
            len(order),
        ),
    )


def _join_items(value: Any) -> str:
    if isinstance(value, list):
        return "；".join(_safe_text(item) for item in value) if value else "无"
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    if value in (None, ""):
        return "无"
    return _safe_text(value)


def _label_with_raw(value: Any, labels: dict[str, str]) -> str:
    raw = _safe_text(value or "unknown")
    return f"{labels.get(raw, raw)}（{raw}）"
