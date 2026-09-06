from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Iterable

from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell


BODY_FONT = "Microsoft YaHei"
SERIF_FONT = "SimSun"
ACCENT = "1F4E79"
HEADER_FILL = "E8EEF7"
LIGHT_FILL = "F6F8FB"
DISCLAIMER = "本报告只表达复现风险、结果差异与人工复核建议，不直接判定论文造假。"
_INLINE_BOLD_PATTERN = re.compile(
    r"(?<!\\)(?:\*\*(?P<asterisk>.+?)\*\*|__(?P<underscore>.+?)__)"
)


def _setup_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    _set_style_font(normal, BODY_FONT, Pt(11))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in (
        ("Heading 1", 15),
        ("Heading 2", 12.5),
        ("Heading 3", 11.5),
    ):
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        _set_style_font(style, BODY_FONT, Pt(size))
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)


def _set_style_font(style: Any, font_name: str, size: Pt) -> None:
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    style.font.size = size


def _set_run_font(
    run: Any,
    font_name: str = BODY_FONT,
    size: Pt | None = None,
) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _add_title(document: DocumentObject, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT)
    _set_run_font(run, BODY_FONT, Pt(20))

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(14)
    run = sub.add_run(subtitle)
    run.font.color.rgb = RGBColor(95, 95, 95)
    _set_run_font(run, SERIF_FONT, Pt(10.5))


def _add_heading(document: DocumentObject, text: str, level: int) -> None:
    paragraph = document.add_heading("", level=level)
    _add_markdown_inline_runs(paragraph, text)


def _add_note(document: DocumentObject, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    _set_run_font(run, BODY_FONT, Pt(10.5))


def _add_appendix_note(
    document: DocumentObject,
    text: str,
    *,
    bold: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.05
    _add_markdown_inline_runs(
        paragraph,
        text,
        bold=bold,
        color=RGBColor(70, 70, 70),
        size=Pt(9.2),
    )


def _add_disclaimer(document: DocumentObject) -> None:
    _add_heading(document, "声明", 1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(DISCLAIMER)
    run.bold = True
    _set_run_font(run, BODY_FONT, Pt(10.5))


def _add_kv_table(
    document: DocumentObject,
    rows: Iterable[tuple[Any, Any]],
) -> None:
    _add_table(document, ["项目", "内容"], [[key, value] for key, value in rows])


def _add_table(
    document: DocumentObject,
    headers: list[str],
    rows: list[list[Any]],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_width(table, 9360)

    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        _set_cell_text(cell, header, bold=True)
        _shade_cell(cell, HEADER_FILL)
    _repeat_table_header(table.rows[0])

    if not rows:
        rows = [["未列出" for _ in headers]]

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            _set_cell_text(cell, _safe_text(value))
            if row_index % 2 == 1:
                _shade_cell(cell, LIGHT_FILL)

    document.add_paragraph()


def _repeat_table_header(row: Any) -> None:
    """Repeat the first row without orphaning it at the bottom of a page."""

    properties = row._tr.get_or_add_trPr()
    marker = properties.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        properties.append(marker)
    marker.set(qn("w:val"), "true")

    # Word can otherwise leave the original header row at the bottom of one
    # page while moving the first data row to the next page.  Keeping every
    # header-cell paragraph with its successor prevents that orphaned start
    # and does not interfere with the repeated header on later pages.
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True


def _set_table_width(table: Any, width_dxa: int) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def _set_table_column_widths(table: Any, widths_dxa: list[int]) -> None:
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            _set_cell_width(cell, width)


def _set_cell_width(cell: _Cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_text(cell: _Cell, text: Any, *, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    _add_markdown_inline_runs(
        paragraph,
        _safe_text(text),
        bold=bold,
        size=Pt(9.5),
    )
    _set_cell_margins(cell, top=90, start=120, bottom=90, end=120)


def _shade_cell(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(
    cell: _Cell,
    top: int,
    start: int,
    bottom: int,
    end: int,
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _add_bullets(document: DocumentObject, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        _add_markdown_inline_runs(
            paragraph,
            _safe_text(item),
            size=Pt(10.5),
        )


def _add_labelled_bullets(
    document: DocumentObject,
    label: str,
    items: Any,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(label)
    run.bold = True
    _set_run_font(run, BODY_FONT, Pt(10.5))
    _add_bullets(document, _items_or_default(items, "未列出"))


def _items_or_default(value: Any, default: str) -> list[str]:
    if isinstance(value, list):
        items = [_safe_text(item) for item in value if _safe_text(item).strip()]
        return items or [default]
    if value:
        return [_safe_text(value)]
    return [default]


def _safe_text(value: Any, limit: int = 1200) -> str:
    if isinstance(value, (dict, list)):
        text = json.dumps(value, ensure_ascii=False, sort_keys=True)
    else:
        text = str(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    if len(text) > limit:
        return text[: limit - 20] + "...[已截断]"
    return text


def _clean_markdown_inline(text: str) -> str:
    cleaned = _safe_text(text)
    cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__([^_]+)__", r"\1", cleaned)
    cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
    return cleaned


def _add_markdown_inline_runs(
    paragraph: Any,
    text: Any,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    font_name: str = BODY_FONT,
    size: Pt | None = None,
) -> None:
    """Append common Markdown inline emphasis as real Word runs.

    The report pipeline currently needs deterministic support for inline code,
    italic cleanup, and the two standard bold spellings.  Keeping this at the
    shared run-construction layer makes paragraphs, lists, and table cells use
    the same behavior instead of leaking Markdown markers into Word.
    """

    raw = _safe_text(text)
    def append_run(
        value: str,
        *,
        inline_bold: bool = False,
        clean: bool = True,
    ) -> None:
        cleaned = _clean_markdown_inline(value) if clean else _safe_text(value)
        if not cleaned:
            return
        run = paragraph.add_run(cleaned)
        if bold or inline_bold:
            run.bold = True
        if italic:
            run.italic = True
        if color is not None:
            run.font.color.rgb = color
        _set_run_font(run, font_name, size)

    for piece in re.split(r"(`[^`]*`)", raw):
        if len(piece) >= 2 and piece.startswith("`") and piece.endswith("`"):
            append_run(piece[1:-1], clean=False)
            continue
        cursor = 0
        for match in _INLINE_BOLD_PATTERN.finditer(piece):
            append_run(piece[cursor : match.start()])
            emphasized = match.group("asterisk")
            if emphasized is None:
                emphasized = match.group("underscore") or ""
            append_run(emphasized, inline_bold=True)
            cursor = match.end()
        append_run(piece[cursor:])

    if not paragraph.runs:
        append_run(raw)


def _save(document: DocumentObject, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path
