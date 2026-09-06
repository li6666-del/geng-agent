from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx.document import Document as DocumentObject

from .docx_assets import (
    _add_image_comparison_table,
    _add_markdown_image,
    _parse_markdown_image_cell,
)
from .docx_styles import (
    _add_appendix_note,
    _add_bullets,
    _add_heading,
    _add_markdown_inline_runs,
    _add_table,
)


def _add_markdown_body(
    document: DocumentObject,
    markdown_text: str,
    *,
    base_dir: Path | None = None,
) -> None:
    lines = markdown_text.splitlines()
    index = 0
    in_appendix = False
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            index += 1
            continue
        table = _parse_markdown_image_table(lines, index)
        if table:
            _add_image_comparison_table(
                document,
                table["headers"],
                table["rows"],
                base_dir=base_dir,
            )
            index += table["consumed"]
            continue
        table = _parse_markdown_table(lines, index)
        if table:
            _add_table(document, table["headers"], table["rows"])
            index += table["consumed"]
            continue
        image_match = re.fullmatch(r"!\[([^\]]*)\]\((.*)\)", line)
        if image_match:
            _add_markdown_image(
                document,
                image_match.group(1).strip(),
                image_match.group(2).strip(),
                base_dir=base_dir,
            )
        elif line.startswith("### "):
            _add_heading(
                document,
                line[4:].strip(),
                3 if not in_appendix else 3,
            )
        elif line.startswith("## "):
            heading = line[3:].strip()
            if heading.startswith("附录"):
                document.add_page_break()
                in_appendix = True
                _add_heading(document, heading, 1)
            else:
                _add_heading(document, heading, 2 if not in_appendix else 3)
        elif line.startswith("# "):
            if in_appendix:
                _add_appendix_note(document, line[2:].strip(), bold=True)
            else:
                _add_heading(document, line[2:].strip(), 1)
        elif line.startswith(("- ", "* ")):
            _add_bullets(document, [line[2:].strip()])
        elif line.startswith("```"):
            pass
        else:
            if in_appendix:
                _add_appendix_note(document, line)
            else:
                paragraph = document.add_paragraph()
                _add_markdown_inline_runs(paragraph, line)
        index += 1


def _parse_markdown_table(
    lines: list[str],
    start: int,
) -> dict[str, Any] | None:
    """Parse a standard pipe-delimited Markdown table of any width."""

    if start + 1 >= len(lines):
        return None
    header = lines[start].strip()
    separator = lines[start + 1].strip()
    if "|" not in header or "|" not in separator:
        return None
    headers = _split_markdown_table_row(header)
    separators = _split_markdown_table_row(separator)
    if not headers or len(headers) != len(separators):
        return None
    if not all(
        re.fullmatch(r":?-{3,}:?", item.strip())
        for item in separators
    ):
        return None

    rows: list[list[str]] = []
    index = start + 2
    while index < len(lines):
        row = lines[index].strip()
        if not row or "|" not in row:
            break
        cells = _split_markdown_table_row(row)
        if len(cells) != len(headers):
            break
        rows.append(cells)
        index += 1
    return {
        "headers": headers,
        "rows": rows,
        "consumed": index - start,
    }


def _parse_markdown_image_table(
    lines: list[str],
    start: int,
) -> dict[str, Any] | None:
    table = _parse_markdown_table(lines, start)
    if table is None or not table["rows"]:
        return None
    if not any(
        _parse_markdown_image_cell(cell) is not None
        for row in table["rows"]
        for cell in row
    ):
        return None
    return table


def _split_markdown_table_row(row: str) -> list[str]:
    text = row.strip()
    if text.startswith("|"):
        text = text[1:]
    if _ends_with_unescaped_pipe(text):
        text = text[:-1]

    cells: list[str] = []
    current: list[str] = []
    in_code = False
    index = 0
    while index < len(text):
        character = text[index]
        if character == "\\" and index + 1 < len(text) and text[index + 1] == "|":
            current.append("|")
            index += 2
            continue
        if character == "`":
            in_code = not in_code
            current.append(character)
        elif character == "|" and not in_code:
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(character)
        index += 1
    cells.append("".join(current).strip())
    return cells


def _ends_with_unescaped_pipe(text: str) -> bool:
    if not text.endswith("|"):
        return False
    backslashes = 0
    for character in reversed(text[:-1]):
        if character != "\\":
            break
        backslashes += 1
    return backslashes % 2 == 0
