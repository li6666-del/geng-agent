from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import _Cell

from .docx_assets import (
    RESULT_REVIEW_COMPARISON_IMAGE_WIDTH_IN,
    RESULT_REVIEW_IMAGE_WIDTH_IN,
    _add_cell_paragraph,
    _add_image_comparison_table,
    _add_markdown_image,
    _add_markdown_image_to_cell,
    _clear_cell,
    _parse_markdown_image_cell,
    _resolve_markdown_image_path,
)
from .docx_markdown import (
    _add_markdown_body,
    _parse_markdown_image_table,
    _split_markdown_table_row,
)
from .docx_reports import (
    ALIGNMENT_LABELS,
    DIMENSION_RATING_LABELS,
    RESULT_REVIEW_DIMENSION_LABELS,
    RISK_LABELS,
    SCIENTIFIC_VERDICT_LABELS,
    _join_items,
    _join_named_files,
    _label_with_raw,
    _ordered_dimension_reviews,
    _runtime_requirement_warning_count,
    _runtime_requirement_warnings,
    populate_result_review_document,
    populate_review_document,
)
from .docx_styles import (
    ACCENT,
    BODY_FONT,
    DISCLAIMER,
    HEADER_FILL,
    LIGHT_FILL,
    SERIF_FONT,
    _add_appendix_note,
    _add_bullets,
    _add_disclaimer,
    _add_heading,
    _add_kv_table,
    _add_labelled_bullets,
    _add_note,
    _add_table,
    _add_title,
    _clean_markdown_inline,
    _items_or_default,
    _safe_text,
    _save,
    _set_cell_margins,
    _set_cell_text,
    _set_cell_width,
    _set_run_font,
    _set_style_font,
    _set_table_column_widths,
    _set_table_width,
    _setup_document,
    _shade_cell,
)


def write_review_docx(
    path: Path,
    *,
    paper: dict[str, Any],
    facts: dict[str, Any],
    tasks: dict[str, Any],
    risk_report: dict[str, Any],
    validation: dict[str, Any],
    runtime_result: dict[str, Any],
    result_review_result: dict[str, Any],
    repro_project_dir: Path,
) -> Path:
    """Create the main engineering reproducibility review report."""

    document = Document()
    _setup_document(document)
    _add_title(
        document,
        "耿同学agent 论文工程复现审查报告",
        "面向通信论文的工程事实抽取、复现代码生成与复现风险评估",
    )
    populate_review_document(
        document,
        paper=paper,
        facts=facts,
        tasks=tasks,
        risk_report=risk_report,
        validation=validation,
        runtime_result=runtime_result,
        result_review_result=result_review_result,
        repro_project_dir=repro_project_dir,
    )
    _add_disclaimer(document)
    return _save(document, path)


def write_result_review_docx(
    path: Path,
    *,
    result_review: dict[str, Any],
    status: dict[str, Any] | None = None,
) -> Path:
    """Create the result-level multimodal review report."""

    document = Document()
    _setup_document(document)
    _add_title(
        document,
        "复现结果二次审查报告",
        "基于本地复现数据、图像与原论文页面图的结果级差异分析",
    )
    populate_result_review_document(
        document,
        result_review=result_review,
        status=status,
    )
    _add_disclaimer(document)
    return _save(document, path)


def write_result_review_markdown_docx(
    path: Path,
    *,
    markdown_text: str,
    status: dict[str, Any] | None = None,
) -> Path:
    """Create the human-readable result review Word report from Markdown."""

    return write_markdown_report_docx(
        path,
        markdown_text=markdown_text,
        title="复现结果二次审查报告",
        subtitle="本地复现结果与论文证据的人工阅读版对比报告",
        base_dir=path.parent,
    )


def write_markdown_report_docx(
    path: Path,
    *,
    markdown_text: str,
    title: str,
    subtitle: str,
    base_dir: Path | None = None,
) -> Path:
    """Render a Codex-authored Markdown report without rewriting its content."""

    document = Document()
    _setup_document(document)
    _add_title(document, title, subtitle)
    _add_markdown_body(document, markdown_text, base_dir=base_dir)
    _add_disclaimer(document)
    return _save(document, path)
