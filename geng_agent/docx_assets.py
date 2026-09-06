from __future__ import annotations

import re
from pathlib import Path

from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell

from .docx_styles import (
    BODY_FONT,
    HEADER_FILL,
    _add_markdown_inline_runs,
    _add_note,
    _safe_text,
    _set_cell_margins,
    _set_cell_text,
    _repeat_table_header,
    _set_run_font,
    _set_table_column_widths,
    _set_table_width,
    _shade_cell,
)


RESULT_REVIEW_IMAGE_WIDTH_IN = 6.2
RESULT_REVIEW_COMPARISON_IMAGE_WIDTH_IN = 3.0


def _add_image_comparison_table(
    document: DocumentObject,
    headers: list[str],
    rows: list[list[str]],
    *,
    base_dir: Path | None = None,
) -> None:
    column_count = len(headers)
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_width(table, 9360)
    base_width, remainder = divmod(9360, column_count)
    column_widths = [base_width] * column_count
    column_widths[-1] += remainder
    _set_table_column_widths(table, column_widths)
    image_width_in = max(
        0.35,
        min(
            RESULT_REVIEW_COMPARISON_IMAGE_WIDTH_IN,
            (base_width / 1440) - 0.25,
        ),
    )

    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, _safe_text(header), bold=True)
        _shade_cell(cell, HEADER_FILL)
    _repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            _clear_cell(cell)
            _set_cell_margins(
                cell,
                top=120,
                start=120,
                bottom=120,
                end=120,
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _add_markdown_image_to_cell(
                cell,
                value,
                base_dir=base_dir,
                image_width_in=image_width_in,
            )
    document.add_paragraph()


def _add_markdown_image_to_cell(
    cell: _Cell,
    value: str,
    *,
    base_dir: Path | None = None,
    image_width_in: float = RESULT_REVIEW_COMPARISON_IMAGE_WIDTH_IN,
) -> None:
    image_items = _parse_markdown_image_cell(value)
    if image_items is None:
        _add_cell_paragraph(
            cell,
            value or "无可用图片",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        return

    for index, (caption, raw_path) in enumerate(image_items):
        image_path = _resolve_markdown_image_path(raw_path, base_dir)
        if not image_path.exists():
            _add_cell_paragraph(
                cell,
                f"图片缺失：{raw_path}",
                align=WD_ALIGN_PARAGRAPH.CENTER,
                italic=True,
            )
            continue
        try:
            paragraph = (
                cell.paragraphs[0]
                if index == 0
                else cell.add_paragraph()
            )
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run()
            run.add_picture(
                str(image_path),
                width=Inches(image_width_in),
            )
            if caption:
                caption_paragraph = cell.add_paragraph()
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_paragraph.paragraph_format.space_after = Pt(4)
                caption_run = caption_paragraph.add_run(_safe_text(caption))
                caption_run.italic = True
                caption_run.font.color.rgb = RGBColor(95, 95, 95)
                _set_run_font(caption_run, BODY_FONT, Pt(9))
        except Exception as exc:
            _add_cell_paragraph(
                cell,
                f"图片插入失败：{raw_path}（{type(exc).__name__}: {exc}）",
                align=WD_ALIGN_PARAGRAPH.CENTER,
                italic=True,
            )


def _parse_markdown_image_cell(value: str) -> list[tuple[str, str]] | None:
    parts = [
        part.strip()
        for part in re.split(
            r"<br\s*/?>",
            value.strip(),
            flags=re.IGNORECASE,
        )
    ]
    if not parts or any(not part for part in parts):
        return None

    images: list[tuple[str, str]] = []
    for part in parts:
        image_match = re.fullmatch(r"!\[([^\]]*)\]\((.*)\)", part)
        if not image_match:
            return None
        images.append(
            (
                image_match.group(1).strip(),
                image_match.group(2).strip(),
            )
        )
    return images


def _add_cell_paragraph(
    cell: _Cell,
    text: str,
    *,
    align: int | None = None,
    italic: bool = False,
    muted: bool = False,
) -> None:
    paragraph = (
        cell.add_paragraph()
        if cell.paragraphs and cell.paragraphs[0].text
        else cell.paragraphs[0]
    )
    paragraph.paragraph_format.space_after = Pt(2)
    if align is not None:
        paragraph.alignment = align
    _add_markdown_inline_runs(
        paragraph,
        _safe_text(text),
        italic=italic,
        color=RGBColor(95, 95, 95) if muted else None,
        size=Pt(9),
    )


def _clear_cell(cell: _Cell) -> None:
    cell.text = ""


def _add_markdown_image(
    document: DocumentObject,
    caption: str,
    raw_path: str,
    *,
    base_dir: Path | None = None,
) -> None:
    image_path = _resolve_markdown_image_path(raw_path, base_dir)
    if not image_path.exists():
        _add_note(document, f"图片缺失：{raw_path}")
        return
    try:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(
            str(image_path),
            width=Inches(RESULT_REVIEW_IMAGE_WIDTH_IN),
        )
        if caption:
            caption_paragraph = document.add_paragraph()
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_paragraph.add_run(_safe_text(caption))
            caption_run.italic = True
            caption_run.font.color.rgb = RGBColor(95, 95, 95)
            _set_run_font(caption_run, BODY_FONT, Pt(9))
    except Exception as exc:
        _add_note(
            document,
            f"图片插入失败：{raw_path}（{type(exc).__name__}: {exc}）",
        )


def _resolve_markdown_image_path(
    raw_path: str,
    base_dir: Path | None,
) -> Path:
    path = Path(raw_path.strip().strip("<>"))
    if path.is_absolute() or base_dir is None:
        return path
    return base_dir / path
